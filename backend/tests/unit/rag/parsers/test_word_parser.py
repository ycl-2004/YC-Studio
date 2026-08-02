"""Tests for ordered DOCX paragraph and table extraction."""

from docx import Document

from app.rag.parsers import WordParser


def test_word_parser_keeps_document_order_and_table_content(tmp_path) -> None:
    source = tmp_path / "report.docx"
    document = Document()
    document.add_heading("季度报告", level=0)
    document.add_paragraph("表格之前")
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "指标"
    table.rows[0].cells[1].text = "数值"
    table.rows[1].cells[0].text = "收入"
    table.rows[1].cells[1].text = "100"
    document.add_paragraph("表格之后")
    document.save(source)

    result = WordParser().parse(source)

    assert result.title == "季度报告"
    assert result.meta["table_count"] == 1
    assert "| 指标 | 数值 |" in result.text
    assert "| 收入 | 100 |" in result.text
    assert result.text.index("表格之前") < result.text.index("| 指标")
    assert result.text.index("| 收入") < result.text.index("表格之后")
