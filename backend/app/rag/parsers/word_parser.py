"""DOCX parser preserving paragraph/table order and table contents."""

from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from docx.table import Table
from docx.text.paragraph import Paragraph

from app.rag.parsers.base import BaseParser, ParseResult
from app.rag.parsers.table_parser import rows_to_markdown


def _table_to_markdown(table: Table) -> str:
    rows = ([cell.text for cell in row.cells] for row in table.rows)
    return rows_to_markdown(rows)


def _document_blocks(document: DocumentObject) -> list[str]:
    blocks: list[str] = []
    # python-docx 1.2 preserves paragraph/table order through iter_inner_content().
    # Source: https://python-docx.readthedocs.io/en/latest/api/document.html
    for item in document.iter_inner_content():
        if isinstance(item, Paragraph):
            if text := item.text.strip():
                blocks.append(text)
        elif isinstance(item, Table):
            if table_markdown := _table_to_markdown(item):
                blocks.append(table_markdown)
    return blocks


class WordParser(BaseParser):
    """Parse Word 2007+ DOCX files and retain tables as Markdown."""

    supported_suffixes = frozenset({".docx"})

    def parse(self, file_path: Path) -> ParseResult:
        document = Document(str(file_path))
        blocks = _document_blocks(document)
        title = next(
            (paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()),
            file_path.stem,
        )
        return ParseResult(
            text="\n\n".join(blocks),
            title=title,
            meta={
                "source_path": str(file_path),
                "table_count": len(document.tables),
            },
            parser_version="docx-python-docx-v1",
        )
